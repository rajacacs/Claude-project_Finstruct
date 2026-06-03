"""DOCX export — Directors Report + Audit Report with placeholder fill."""

from __future__ import annotations
from pathlib import Path
from datetime import datetime
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH


_BLUE  = RGBColor(0, 120, 212)
_DARK  = RGBColor(16, 110, 190)
_TEXT  = RGBColor(32, 31, 30)


def _h(doc: Document, text: str, level: int = 1):
    p = doc.add_heading(text, level=level)
    p.runs[0].font.color.rgb = _BLUE
    p.runs[0].font.name = "Segoe UI"


def _para(doc: Document, text: str, bold: bool = False, italic: bool = False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold   = bold
    run.italic = italic
    run.font.name = "Segoe UI"
    run.font.size = Pt(10)
    run.font.color.rgb = _TEXT
    return p


OPINION_PARAGRAPHS = {
    "Unmodified": (
        "In our opinion and to the best of our information and according to the explanations "
        "given to us, the aforesaid financial statements give the information required by the "
        "Companies Act, 2013 in the manner so required and give a true and fair view in conformity "
        "with the accounting principles generally accepted in India of the state of affairs of the "
        "Company as at 31st March {{FY_END_YEAR}}, and of its profit / (loss) and cash flows for "
        "the year ended on that date."
    ),
    "Qualified": (
        "In our opinion and to the best of our information and according to the explanations given "
        "to us, **except for the effects of the matters described in the Basis for Qualified "
        "Opinion paragraph above**, the aforesaid financial statements give the information "
        "required by the Companies Act, 2013 in the manner so required and give a true and fair "
        "view in conformity with the accounting principles generally accepted in India of the "
        "state of affairs of the Company as at 31st March {{FY_END_YEAR}}, and of its profit / (loss) "
        "and cash flows for the year ended on that date."
    ),
    "Adverse": (
        "In our opinion, **because of the significance of the matters described in the Basis for "
        "Adverse Opinion paragraph above**, the aforesaid financial statements do NOT give a true "
        "and fair view in conformity with the accounting principles generally accepted in India "
        "of the state of affairs of the Company as at 31st March {{FY_END_YEAR}}, or of its profit / "
        "(loss) and cash flows for the year ended on that date."
    ),
    "Disclaimer": (
        "**We do not express an opinion** on the accompanying financial statements of the Company. "
        "Because of the significance of the matters described in the Basis for Disclaimer of "
        "Opinion paragraph above, we have not been able to obtain sufficient appropriate audit "
        "evidence to provide a basis for an audit opinion on these financial statements."
    ),
}


def _fill(template: str, em: dict, extras: dict | None = None) -> str:
    fy = em.get("financial_year", "")
    # Derive FY end year: "2024-25" → "2025"
    fy_end = ""
    if fy and "-" in fy:
        parts = fy.split("-")
        try:
            start = int(parts[0])
            fy_end = str(start + 1)
        except ValueError:
            fy_end = parts[-1]
    opinion = em.get("opinion_type", "Unmodified")
    opinion_para = OPINION_PARAGRAPHS.get(opinion, OPINION_PARAGRAPHS["Unmodified"])

    subs = {
        "{{COMPANY_NAME}}":   em.get("entity_name", em.get("Company_Name", "[Company]")),
        "{{CIN}}":            em.get("cin", em.get("CIN", "")),
        "{{ADDRESS}}":        em.get("address", em.get("Registered_Office", "")),
        "{{FY}}":             fy,
        "{{FY_END_YEAR}}":    fy_end,
        "{{OPINION_TYPE}}":   opinion,
        "{{OPINION_PARA}}":   opinion_para,
        "{{UDIN}}":           em.get("udin", ""),
        "{{SIGN_DATE}}":      em.get("signing_date", em.get("Signing_Date",
                                     datetime.now().strftime("%d %B %Y"))),
        "{{SIGN_PLACE}}":     em.get("signing_place", em.get("Signing_Place", "")),
        "{{DIR_1}}":          em.get("dir1_name", em.get("DIR_1_NAME", "[Director 1]")),
        "{{DIR_1_DESIG}}":    em.get("dir1_desig", em.get("DIR_1_DESIG", "Director")),
        "{{DIR_1_DIN}}":      em.get("dir1_din", em.get("DIR_1_DIN", "")),
        "{{DIR_2}}":          em.get("dir2_name", em.get("DIR_2_NAME", "")),
        "{{AUDITOR_FIRM}}":   em.get("auditor_firm", em.get("Auditor_Firm", "[Auditor Firm]")),
        "{{AUDITOR_FRN}}":    em.get("auditor_frn", em.get("Auditor_Firm_Reg", "")),
        "{{AUDITOR_PARTNER}}":em.get("auditor_partner", em.get("Auditor_Partner", "")),
        "{{AUDITOR_MRN}}":    em.get("auditor_mrn", em.get("Auditor_MemNo", "")),
    }
    if extras:
        subs.update(extras)
    for k, v in subs.items():
        template = template.replace(k, v or "")
    # Replace OPINION section AFTER substitution so opinion paragraph also gets {{FY_END_YEAR}} subbed
    template = template.replace("{{FY_END_YEAR}}", fy_end)
    return template


DIRECTORS_REPORT_TEMPLATE = """DIRECTORS' REPORT

To,
The Members,
{{COMPANY_NAME}}

Your Directors present the Annual Report of {{COMPANY_NAME}} (CIN: {{CIN}}) for the Financial Year ended 31st March.

1. FINANCIAL RESULTS
The highlights of the financial performance of the Company for the Financial Year {{FY}} are as follows:
[Insert financial summary table here]

2. OPERATIONS
During the year under review, the Company carried on its business activities. [Describe key operations]

3. DIVIDEND
The Board of Directors does not recommend any dividend for the Financial Year {{FY}}. / The Board of Directors recommends a dividend of ₹__ per share.

4. RESERVES
No amount has been proposed to be transferred to Reserves.

5. DIRECTORS
The following are the Directors of the Company:
a) {{DIR_1}} — {{DIR_1_DESIG}} (DIN: {{DIR_1_DIN}})
b) {{DIR_2}} — Director

6. DIRECTORS' RESPONSIBILITY STATEMENT
Pursuant to Section 134(3)(c) of the Companies Act, 2013, the Directors confirm that:
(i) in the preparation of the Annual Accounts, the applicable accounting standards have been followed;
(ii) the accounting policies selected and applied are consistent and judgements made are reasonable;
(iii) proper and sufficient care has been taken for maintenance of adequate accounting records;
(iv) the Annual Accounts have been prepared on a going concern basis.

7. AUDITORS
M/s {{AUDITOR_FIRM}} (FRN: {{AUDITOR_FRN}}), Chartered Accountants, were appointed as Statutory Auditors.

8. SECRETARIAL AUDIT
[If applicable — insert secretarial audit observations]

9. RELATED PARTY TRANSACTIONS
All related party transactions that were entered during FY {{FY}} were on an arm's length basis.

10. ACKNOWLEDGEMENT
The Directors wish to express their gratitude to clients, bankers, employees, and stakeholders.

For and on behalf of the Board of Directors
{{COMPANY_NAME}}

{{DIR_1}}                    {{DIR_2}}
{{DIR_1_DESIG}}               Director
DIN: {{DIR_1_DIN}}

Place: {{SIGN_PLACE}}
Date: {{SIGN_DATE}}
"""

AUDIT_REPORT_TEMPLATE = """INDEPENDENT AUDITOR'S REPORT

To,
The Members,
{{COMPANY_NAME}}

REPORT ON THE AUDIT OF THE FINANCIAL STATEMENTS

OPINION ({{OPINION_TYPE}})
We have audited the accompanying financial statements of {{COMPANY_NAME}} (CIN: {{CIN}}), which comprise the Balance Sheet as at 31st March {{FY_END_YEAR}}, the Statement of Profit and Loss and the Cash Flow Statement for the year then ended, and notes to the financial statements, including a summary of significant accounting policies.

{{OPINION_PARA}}

BASIS FOR OPINION
We conducted our audit in accordance with the Standards on Auditing (SAs) specified under section 143(10) of the Companies Act, 2013. Our responsibilities under those Standards are further described in the Auditor's Responsibilities for the Audit of the Financial Statements section of our report. We are independent of the Company in accordance with the Code of Ethics issued by the Institute of Chartered Accountants of India, and we have fulfilled our other ethical responsibilities in accordance with these requirements. We believe that the audit evidence we have obtained is sufficient and appropriate to provide a basis for our opinion.

KEY AUDIT MATTERS
[Insert key audit matters if applicable]

INFORMATION OTHER THAN THE FINANCIAL STATEMENTS AND AUDITOR'S REPORT THEREON
The Company's Board of Directors is responsible for the other information. [Continue as applicable]

RESPONSIBILITIES OF MANAGEMENT AND THOSE CHARGED WITH GOVERNANCE
The Company's Board of Directors is responsible for the matters stated in section 134(5) of the Companies Act, 2013.

AUDITOR'S RESPONSIBILITIES FOR THE AUDIT OF THE FINANCIAL STATEMENTS
Our objectives are to obtain reasonable assurance about whether the financial statements as a whole are free from material misstatement.

REPORT ON OTHER LEGAL AND REGULATORY REQUIREMENTS
1. As required by the Companies (Auditor's Report) Order, 2020 ("the Order"), issued by the Central Government, we give in the Annexure A, a statement on the matters specified in paragraphs 3 and 4 of the Order.
2. As required by Section 143(3) of the Act, we report that:
   (a) We have sought and obtained all the information and explanations required.
   (b) In our opinion, proper books of account have been kept.
   (c) The Balance Sheet, Statement of Profit and Loss, and Cash Flow Statement are in agreement with the books of account.
   (d) In our opinion, the financial statements comply with the applicable Accounting Standards.

For {{AUDITOR_FIRM}}
Chartered Accountants
FRN: {{AUDITOR_FRN}}

{{AUDITOR_PARTNER}}
Partner
Membership No.: {{AUDITOR_MRN}}
UDIN: {{UDIN}}

Place: {{SIGN_PLACE}}
Date: {{SIGN_DATE}}
"""


def export_docx(entity_master: dict, output_path: Path,
                directors_report_text: str | None = None,
                audit_report_text: str | None = None):
    doc = Document()

    # Set margins
    for section in doc.sections:
        section.left_margin  = Inches(1.0)
        section.right_margin = Inches(1.0)
        section.top_margin   = Inches(1.0)
        section.bottom_margin= Inches(1.0)

    dr_text = directors_report_text or _fill(DIRECTORS_REPORT_TEMPLATE, entity_master)
    ar_text = audit_report_text      or _fill(AUDIT_REPORT_TEMPLATE,    entity_master)

    # Directors Report
    _h(doc, "DIRECTORS' REPORT", 1)
    for line in dr_text.split("\n")[1:]:  # skip first line (already the heading)
        if line.strip().isupper() and len(line.strip()) > 3:
            _h(doc, line.strip(), 2)
        else:
            _para(doc, line)

    doc.add_page_break()

    # Audit Report
    _h(doc, "INDEPENDENT AUDITOR'S REPORT", 1)
    for line in ar_text.split("\n")[1:]:
        if line.strip().isupper() and len(line.strip()) > 3:
            _h(doc, line.strip(), 2)
        else:
            _para(doc, line)

    doc.save(output_path)
